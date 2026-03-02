# -*- coding: utf-8 -*-
"""
Hot News Ranker (CN) — static site generator
- Max 3 items per source; cross-site de-dup; per-item summary
- Outputs: index.html (static), optional TXT/Word
- Default output dir (Windows): D:\MyCode\Hot_Points (override via CLI --outdir)
"""
import os, re, time, argparse, html as htmllib
import datetime as dt
from datetime import timezone
from typing import List, Dict, Tuple, Optional, Any

import requests
from bs4 import BeautifulSoup
import feedparser

DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    DOCX_AVAILABLE = False

OUTPUT_DIR = os.environ.get("HOTNEWS_OUTDIR", r"D:\MyCode\Hot_Points")
os.makedirs(OUTPUT_DIR, exist_ok=True)

HEADERS = {"User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) "
                           "Chrome/122.0 Safari/537.36")}
TIMEOUT=8; RETRY=2; SLEEP_BETWEEN=0.7
MAX_ITEMS_PER_SOURCE=3; TOP_N=80

EXCLUDE_KEYWORDS=["习近平","总书记","国家主席","中共中央","中央政治局","中央委员会",
                  "中央政府","国务院常务会议","中央纪委","中央统战部","国家领导人"]
EXCLUDE_REGEX=re.compile("|".join(map(re.escape,EXCLUDE_KEYWORDS)),re.IGNORECASE)
HOT_KEYWORDS={"突发":3,"通报":2,"最新":2,"预警":2,"发布":1,"春运":3,"消费":2,
              "房产":2,"楼市":2,"经济":2,"事故":3,"暴雪":2,"寒潮":2,"高铁":2,
              "医保":2,"大模型":2,"AI":1,"新能源":2,"锂电":2,"芯片":2,"文旅":2,
              "免税":1,"通行":1,"航班":1}
CITY_KEYWORDS=["广州","深圳","北京","上海","杭州","南京","天津","重庆","武汉","西安"]
SOURCE_WEIGHT={"央视网-新闻频道":4,"央视网-国内新闻":4,"新华网-首页":4,
               "中新网-即时":4,"中新网-要闻":4,"中新网-国内":4,"中新网-社会":3,
               "环球网-国内":2,"网易新闻-国内":2}
SOURCES=[
 ("中新网-即时",{"rss":"https://www.chinanews.com.cn/rss/scroll-news.xml"}),
 ("中新网-要闻",{"rss":"https://www.chinanews.com.cn/rss/importnews.xml"}),
 ("中新网-国内",{"rss":"https://www.chinanews.com.cn/rss/china.xml"}),
 ("中新网-社会",{"rss":"https://www.chinanews.com.cn/rss/society.xml"}),
 ("央视网-新闻频道",{"html":"https://news.cctv.cn/","parser":"parse_cctv_index"}),
 ("央视网-国内新闻",{"html":"https://news.cctv.com/china/","parser":"parse_cctv_china"}),
 ("新华网-首页",{"html":"https://www.news.cn/","parser":"parse_xinhua"}),
 ("环球网-国内",{"html":"https://china.huanqiu.com/","parser":"parse_huanqiu_china"}),
 ("网易新闻-国内",{"html":"https://news.163.com/domestic/","parser":"parse_net163_domestic"}),
]

def get_html(url:str)->str:
    last=None
    for _ in range(RETRY+1):
        try:
            r=requests.get(url,headers=HEADERS,timeout=TIMEOUT)
            r.raise_for_status(); r.encoding=r.apparent_encoding or 'utf-8'
            return r.text
        except Exception as e:
            last=e; time.sleep(0.5)
    raise RuntimeError(f"请求失败：{url} 错误：{last}")

def normalize_space(s: Optional[str]) -> str:
    return re.sub(r"\s+"," ", s or "").strip()

def safe_url(u:str)->str:
    if not u: return u
    u=u.replace(' ','')
    if u.startswith('//'): u='https:'+u
    return u

def strip_html(raw:str)->str:
    if not raw: return ""
    t=re.sub(r"<script.*?>.*?</script>|<style.*?>.*?</style>","",raw,flags=re.S|re.I)
    t=re.sub(r"<[^>]+>","",t)
    return normalize_space(t)

def smart_trim(s:str,max_len:int=160)->str:
    if not s: return s
    if len(s)<=max_len: return s
    cut=s[:max_len+10]
    m=re.search(r"[。！？；.!?;]\s*\S*$",cut)
    return cut[:m.start()+1] if m else cut[:max_len].rstrip("，,;；、.。")+"…"

def fetch_rss(feed_url:str, source_name:str)->List[Dict[str,Any]]:
    d=feedparser.parse(feed_url); items=[]
    for e in d.entries[:MAX_ITEMS_PER_SOURCE]:
        title=normalize_space(getattr(e,'title',''))
        link=safe_url(getattr(e,'link','') or getattr(e,'id',''))
        if not title or not link: continue
        summary=strip_html(getattr(e,'summary',''))
        published=normalize_space(getattr(e,'published','')) or normalize_space(getattr(e,'updated',''))
        pub_parsed=getattr(e,'published_parsed',None) or getattr(e,'updated_parsed',None)
        items.append({"title":title,"url":link,"summary":summary,"published":published,
                      "published_parsed":pub_parsed,"source":source_name,"via":"rss"})
    return items

from bs4 import BeautifulSoup

def parse_cctv_index(html:str)->List[Dict[str,Any]]:
    soup=BeautifulSoup(html,'lxml'); items=[]
    for a in soup.select('a'):
        title=normalize_space(a.get_text()); href=safe_url(a.get('href') or '')
        if not title or not href or not href.startswith('http'): continue
        if len(title)<6: continue
        if any(seg in href for seg in ['/news.','/202','/china','/society','/economy','/tech','/sports']):
            items.append({'title':title,'url':href})
        if len(items)>=MAX_ITEMS_PER_SOURCE: break
    return items

def parse_cctv_china(html:str)->List[Dict[str,Any]]:
    soup=BeautifulSoup(html,'lxml'); items=[]
    for a in soup.select('a'):
        title=normalize_space(a.get_text()); href=safe_url(a.get('href') or '')
        if not title or not href or not href.startswith('http'): continue
        if len(title)<6: continue
        if '/china/' in href or '/202' in href or '/news' in href:
            items.append({'title':title,'url':href})
        if len(items)>=MAX_ITEMS_PER_SOURCE: break
    return items

def parse_xinhua(html:str)->List[Dict[str,Any]]:
    soup=BeautifulSoup(html,'lxml'); items=[]
    for a in soup.select('a'):
        title=normalize_space(a.get_text()); href=safe_url(a.get('href') or '')
        if not title or not href or not href.startswith('http'): continue
        if any(x in href for x in ['#','javascript:']): continue
        if any(seg in href for seg in ['/202','/local/','/fortune/','/world/','/tech/','/photo/','/sports/']):
            if len(title)>=6: items.append({'title':title,'url':href})
        if len(items)>=MAX_ITEMS_PER_SOURCE: break
    return items

def parse_huanqiu_china(html:str)->List[Dict[str,Any]]:
    soup=BeautifulSoup(html,'lxml'); items=[]
    for a in soup.select('a'):
        title=normalize_space(a.get_text()); href=safe_url(a.get('href') or '')
        if not title or not href or not href.startswith('http'): continue
        if len(title)<6: continue
        if 'huanqiu.com' in href or href.startswith('https://china.huanqiu.com'):
            items.append({'title':title,'url':href})
        if len(items)>=MAX_ITEMS_PER_SOURCE: break
    return items

def parse_net163_domestic(html:str)->List[Dict[str,Any]]:
    soup=BeautifulSoup(html,'lxml'); items=[]
    for a in soup.select('a'):
        title=normalize_space(a.get_text()); href=safe_url(a.get('href') or '')
        if not title or not href or not href.startswith('http'): continue
        if len(title)<6: continue
        if 'news.163.com' in href or href.startswith('https://www.163.com'):
            items.append({'title':title,'url':href})
        if len(items)>=MAX_ITEMS_PER_SOURCE: break
    return items

PARSERS={'parse_cctv_index':parse_cctv_index,'parse_cctv_china':parse_cctv_china,
         'parse_xinhua':parse_xinhua,'parse_huanqiu_china':parse_huanqiu_china,
         'parse_net163_domestic':parse_net163_domestic}

def fetch_from_source(name:str, conf:Dict[str,str])->List[Dict[str,Any]]:
    try:
        if 'rss' in conf: return fetch_rss(conf['rss'], name)
        if 'html' in conf and 'parser' in conf:
            html=get_html(conf['html']); parser=PARSERS[conf['parser']]
            data=parser(html)
            for d in data: d['source']=name; d['via']='html'
            return data
        return []
    except Exception:
        return []

def is_excluded(title:str, summary:str='')->bool:
    return bool(EXCLUDE_REGEX.search(f"{title} {summary}"))

def dedup_and_group(items:List[Dict[str,Any]])->List[Dict[str,Any]]:
    def canon_title(t:str)->str:
        t=re.sub(r"（.*?）|\(.*?\)|\[.*?\]|【.*?】","",t); return normalize_space(t)
    buckets={}
    for it in items:
        key=canon_title(it['title'])
        if key not in buckets:
            buckets[key]={'title':key,'url':it['url'],'summary':it.get('summary',''),
                          'published':it.get('published',''),'published_parsed':it.get('published_parsed'),
                          'sources':set([it.get('source','')]),'via':set([it.get('via','')]),'raw':[it]}
        else:
            b=buckets[key]; b['sources'].add(it.get('source','')); b['via'].add(it.get('via',''))
            if ('news.' in it['url'] or '/202' in it['url']) and 'video' not in it['url']: b['url']=it['url']
            if it.get('published_parsed') and not b.get('published_parsed'):
                b['published_parsed']=it['published_parsed']; b['published']=it.get('published','')
            if it.get('summary') and not b.get('summary'): b['summary']=it['summary']
            b['raw'].append(it)
    merged=[]
    for _,v in buckets.items():
        merged.append({'title':v['title'],'url':v['url'],'summary':v.get('summary',''),
                       'published':v.get('published',''),'published_parsed':v.get('published_parsed'),
                       'sources':sorted([s for s in v['sources'] if s]),
                       'via_list':sorted([s for s in v['via'] if s]),'raw':v['raw']})
    return merged

def fetch_page_summary(url:str)->str:
    try:
        html=get_html(url); soup=BeautifulSoup(html,'lxml')
        for sel in ['meta[name="description"]','meta[name="Description"]',
                    'meta[property="og:description"]','meta[name="og:description"]']:
            tag=soup.select_one(sel)
            if tag and tag.get('content'): return strip_html(tag['content'])[:400]
        for p in soup.select('article p, .content p, .article p, .main p, p'):
            text=normalize_space(p.get_text())
            if len(text)>=30: return text[:400]
    except Exception:
        pass
    return ''

def attach_summaries(items:List[Dict[str,Any]])->None:
    for it in items:
        s=strip_html(it.get('summary',''))
        if not s: s=fetch_page_summary(it['url'])
        if not s: s=it['title']
        it['summary_final']=smart_trim(s,160)

def score_item(it:Dict[str,Any])->float:
    score=0.0; title=it.get('title',''); summary=it.get('summary_final') or it.get('summary','')
    txt=f"{title} {summary}"; src_count=len(it.get('sources',[]))
    if src_count>=3: score+=6
    elif src_count==2: score+=4
    elif src_count==1: score+=1
    srcw=[SOURCE_WEIGHT.get(s,1) for s in it.get('sources',[])]
    if srcw: score+=min(sum(srcw),10)
    for k,w in HOT_KEYWORDS.items():
        if k in txt: score+=w
    if re.search(r"\d",title): score+=1
    if any(c in title for c in CITY_KEYWORDS): score+=1
    if 10<=len(title)<=30: score+=1
    pp=it.get('published_parsed')
    if pp:
        try:
            pub_dt=dt.datetime(*pp[:6],tzinfo=timezone.utc)
            diff_h=(dt.datetime.now(timezone.utc)-pub_dt).total_seconds()/3600.0
            if diff_h<3: score+=3
            elif diff_h<8: score+=2
            elif diff_h<24: score+=1
        except Exception: pass
    return score

def rank_items(items:List[Dict[str,Any]])->List[Dict[str,Any]]:
    return sorted(items,key=score_item,reverse=True)

def fmt_pub_time(it:Dict[str,Any])->str:
    if it.get('published_parsed'):
        try:
            pub_dt=dt.datetime(*it['published_parsed'][:6],tzinfo=timezone.utc).astimezone()
            return pub_dt.strftime('%Y-%m-%d %H:%M')
        except Exception:
            return it.get('published','')
    return it.get('published','')

def save_to_txt(items:List[Dict[str,Any]], out_fullpath:str):
    now=dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    with open(out_fullpath,'w',encoding='utf-8') as f:
        f.write('今日中国热点新闻（自动评分排序，已过滤国家领导人及中央政府相关)')
        f.write(f'生成时间：{now}')
        f.write('数据源：中新网(RSS优先)、央视网、新华网、环球网、网易（部分HTML解析）')
        f.write('评分维度：多源交叉/来源权重/关键词/结构/时间新鲜度')
        for i,it in enumerate(items,1):
            srcs='、'.join(it.get('sources',[])); pub=fmt_pub_time(it)
            f.write(f"{i:02d}. {it['title']}（来源：{srcs or '未知'}；日期：{pub or '—'}）")
            f.write(f"    摘要：{it.get('summary_final','')}")
            f.write(f"    {it['url']}")

def add_hyperlink(paragraph, text, url):
    r_id=paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink=OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    run=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    color=OxmlElement('w:color'); color.set(qn('w:val'),'0000FF'); rPr.append(color)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    run.append(rPr); t=OxmlElement('w:t'); t.text=text; run.append(t)
    hyperlink.append(run); paragraph._p.append(hyperlink)

def save_to_word(items:List[Dict[str,Any]], out_fullpath:str):
    if not DOCX_AVAILABLE:
        print('[提示] 未检测到 python-docx，已跳过 Word 生成。'); return
    doc=Document(); doc.add_heading('今日中国热点新闻（自动评分排序）',level=1)
    for i,it in enumerate(items,1):
        title=it['title']; url=it['url']; src='、'.join(it.get('sources',[])) or '未知'; pub=fmt_pub_time(it) or '—'
        p=doc.add_paragraph(); p.add_run(f'{i:02d}. '); add_hyperlink(p,title,url)
        doc.add_paragraph(f'（来源：{src}；日期：{pub}）')
        doc.add_paragraph(f"摘要：{it.get('summary_final','')}")
        doc.add_paragraph('')
    doc.save(out_fullpath)

def save_to_html(items:List[Dict[str,Any]], out_fullpath:str):
    now=dt.datetime.now().strftime('%Y-%m-%d %H:%M')
    css=(":root{--fg:#222;--muted:#666;--link:#0969da;--bg:#fff;--card:#f8f9fa}"          "*{box-sizing:border-box}"          "body{margin:0;padding:24px 16px;color:var(--fg);background:var(--bg);"          "font:16px/1.6 system-ui,-apple-system,Segoe UI,Roboto,Arial,"Microsoft Yahei",sans-serif}"          ".wrap{max-width:860px;margin:0 auto}"          "header h1{margin:0 0 4px 0;font-size:1.6rem}"          "header .ts{color:var(--muted);font-size:.9rem;margin-bottom:16px}"          ".card{background:var(--card);border:1px solid #e5e7eb;border-radius:10px;padding:14px 16px;margin:10px 0}"          ".idx{display:inline-block;width:36px;color:#888}"          ".title a{color:var(--link);text-decoration:none}"          ".title a:hover{text-decoration:underline}"          ".meta{color:var(--muted);font-size:.9rem;margin-top:4px}"          ".sum{margin-top:6px}"          "footer{color:var(--muted);font-size:.85rem;margin-top:28px}")
    head = (
        '<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"/>'
        '<meta name="viewport" content="width=device-width,initial-scale=1"/>'
        '<title>今日中国热点新闻</title><style>'+css+'</style></head><body>'
        '<div class="wrap"><header><h1>今日中国热点新闻</h1>'
        f'<div class="ts">生成时间：{now}</div></header><section>'
    )
    rows=[]
    for i,it in enumerate(items,1):
        title=htmllib.escape(it['title']); url=htmllib.escape(it['url'])
        srcs='、'.join(it.get('sources',[])) or '未知'; pub=fmt_pub_time(it) or '—'
        summary=htmllib.escape(it.get('summary_final','') or it.get('summary','') or it['title'])
        rows.append(
            f"<article class='card'><div class='title'><span class='idx'>{i:02d}.</span>"
            f"<a href="{url}" target="_blank" rel="noopener noreferrer">{title}</a></div>"
            f"<div class='meta'>来源：{htmllib.escape(srcs)}；日期：{htmllib.escape(pub)}</div>"
            f"<div class='sum'>摘要：{summary}</div></article>"
        )
    tail = "</section><footer>本页由定时任务自动生成（RSS/权威站点抓取 + 评分排序）。</footer></div></body></html>"
    with open(out_fullpath,'w',encoding='utf-8') as f:
        f.write(head + "".join(rows) + tail)

def main():
    parser=argparse.ArgumentParser(description='CN Hot News Ranker')
    parser.add_argument('--html', default=None, help='输出 HTML 文件路径，如 D:\MyCode\Hot_Points\index.html')
    parser.add_argument('--no-txt', action='store_true', help='不生成 TXT')
    parser.add_argument('--no-docx', action='store_true', help='不生成 Word')
    parser.add_argument('--outdir', default=OUTPUT_DIR, help='输出目录（覆盖默认设置）')
    args=parser.parse_args()

    outdir=args.outdir; os.makedirs(outdir, exist_ok=True)
    print('=== 中国热点新闻评分模型（固定输出目录）启动 ===')
    print('输出目录：', outdir)

    raw=[]
    for name, conf in SOURCES:
        time.sleep(SLEEP_BETWEEN)
        print(f'[抓取] {name} ... ', end='')
        data=fetch_from_source(name, conf)
        before=len(data)
        data=[d for d in data if not is_excluded(d.get('title',''), d.get('summary',''))]
        after=len(data)
        print(f'获取 {before} 条，过滤后 {after} 条')
        raw.extend(data)
    print(f'[汇总] 原始合计 {len(raw)} 条；去重合并 …')
    merged=dedup_and_group(raw)
    print('[摘要] 生成摘要 …'); attach_summaries(merged)
    print(f'[评分] 排序 …')
    ranked=rank_items(merged); topn=ranked[:TOP_N]
    today=dt.date.today().strftime('%Y-%m-%d')

    if not args.no_txt:
        txt_path=os.path.join(outdir, f'hot_news_{today}.txt'); save_to_txt(topn, txt_path)
        print('[完成] TXT:', txt_path)
    if not args.no_docx:
        docx_path=os.path.join(outdir, f'hot_news_{today}.docx'); save_to_word(topn, docx_path)
        if DOCX_AVAILABLE: print('[完成] Word:', docx_path)
        else: print('[提示] 未安装 python-docx，已跳过 Word。')
    if args.html:
        save_to_html(topn, args.html); print('[完成] HTML:', args.html)

if __name__=='__main__':
    import traceback
    try:
        main()
    except Exception as e:
        log_path=os.path.join(OUTPUT_DIR,'error.log')
        with open(log_path,'a',encoding='utf-8') as f:
            f.write(f"[{dt.datetime.now()}] {e}{traceback.format_exc()}")
        print('[异常] 运行出错，已写入日志：', log_path)
